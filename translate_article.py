import os
import argparse
import sys
from dotenv import load_dotenv
from pdf2docx import Converter
import docx
import time
from progress import update_progress
# Try to import the correct module for the Gemini API
try:
    from google import genai
    from google.genai import types
except ImportError:
    print("Error: Could not import google.genai. Make sure you installed 'google-genai'.")
    sys.exit(1)

def convert_pdf_to_docx(pdf_path, docx_path):
    print(f"Converting PDF '{pdf_path}' to DOCX '{docx_path}'...")
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        print("Conversion successful.")
    except Exception as e:
        print(f"Error converting PDF to DOCX: {e}")
        sys.exit(1)

def translate_text(client, text, direction="en2es", retries=3):
    if not text or not text.strip():
        return text

    # Define the prompts based on direction
    if direction == "en2es":
        prompt = (
            "Translate the following text from academic English to academic Mexican Spanish. "
            "The context is strictly 'Mathematics Education' (Educación Matemática). "
            "Maintain the academic style, flow, and exact meaning. "
            "Do not add any conversational filler, markdown formatting (unless present in original), or explanations. "
            "Provide ONLY the translated text:\n\n"
        )
    else:
        prompt = (
            "Translate the following text from academic Spanish to academic English. "
            "The context is strictly 'Mathematics Education' (Educación Matemática). "
            "Maintain the academic style, flow, and exact meaning. "
            "Do not add any conversational filler, markdown formatting (unless present in original), or explanations. "
            "Provide ONLY the translated text:\n\n"
        )

    full_prompt = prompt + text

    for attempt in range(retries):
        try:
            # Respect rate limit (15 RPM -> 1 request per 4.1 seconds)
            time.sleep(4.1)
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=full_prompt,
            )
            translated_text = response.text.strip()
            return translated_text
        except Exception as e:
            err_msg = str(e)
            print(f"Warning: Failed to translate text on attempt {attempt+1}: {err_msg[:100]}...")
            if '429' in err_msg or 'RESOURCE_EXHAUSTED' in err_msg or '503' in err_msg or 'UNAVAILABLE' in err_msg:
                # wait a bit longer if we hit rate limits or service unavailable
                time.sleep(20)
            else:
                # Still wait and retry for other transient errors
                time.sleep(10)
                
    return text # return original if failed all retries

def translate_batch_text(client, texts, direction="en2es", retries=3):
    if not texts:
        return []
        
    delimiter = "\n\n|||\n\n"
    combined_text = delimiter.join(texts)
    
    if direction == "en2es":
        prompt = (
            "Translate the following text segments from academic English to academic Mexican Spanish. "
            "The context is strictly 'Mathematics Education' (Educación Matemática). "
            "Maintain the academic style, flow, and exact meaning. "
            "IMPORTANT: The text contains multiple distinct segments separated by '|||'. "
            "You MUST separate your translations with the exact same '|||' delimiter. "
            "Provide ONLY the translated text segments:\n\n"
        )
    else:
        prompt = (
            "Translate the following text segments from academic Spanish to academic English. "
            "The context is strictly 'Mathematics Education' (Educación Matemática). "
            "Maintain the academic style, flow, and exact meaning. "
            "IMPORTANT: The text contains multiple distinct segments separated by '|||'. "
            "You MUST separate your translations with the exact same '|||' delimiter. "
            "Provide ONLY the translated text segments:\n\n"
        )
        
    full_prompt = prompt + combined_text

    for attempt in range(retries):
        try:
            time.sleep(4.1)
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=full_prompt,
            )
            translated_combined = response.text.strip()
            translated_segments = [s.strip() for s in translated_combined.split('|||')]
            return translated_segments
        except Exception as e:
            err_msg = str(e)
            print(f"Warning: Failed to translate batch on attempt {attempt+1}: {err_msg[:100]}...")
            if '429' in err_msg or 'RESOURCE_EXHAUSTED' in err_msg or '503' in err_msg or 'UNAVAILABLE' in err_msg:
                time.sleep(20)
            else:
                time.sleep(10)
                
    return texts

def process_docx(docx_path, output_path, client, direction="en2es", task_id=None, mode="paragraph"):
    print(f"Opening DOCX '{docx_path}' for translation...")
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error opening DOCX: {e}")
        if task_id: update_progress(task_id, "error", 0, f"Error opening DOCX: {e}")
        sys.exit(1)

    print("Translating paragraphs...")
    valid_paras = [p for p in doc.paragraphs if p.text.strip()]
    total_paras = len(valid_paras)
    
    if mode == "batch":
        batch_size = 5
        for i in range(0, total_paras, batch_size):
            chunk = valid_paras[i:i+batch_size]
            texts = [p.text.strip() for p in chunk]
            
            translated_texts = translate_batch_text(client, texts, direction)
            
            if len(translated_texts) == len(chunk):
                for j, para in enumerate(chunk):
                    para.text = translated_texts[j]
            else:
                print(f"Batch mismatch (expected {len(chunk)}, got {len(translated_texts)}). Falling back to singular translation for this batch.")
                for para in chunk:
                    para.text = translate_text(client, para.text.strip(), direction)
                    
            if task_id:
                prog = 40 + int((min(i + batch_size, total_paras) / total_paras) * 50)
                update_progress(task_id, "processing", prog, f"Traduciendo (por lotes) {min(i + batch_size, total_paras)} de {total_paras}...")
    else:
        for i, para in enumerate(valid_paras):
            translated_text = translate_text(client, para.text.strip(), direction)
            para.text = translated_text
            
            if i % 2 == 0 and task_id:
                prog = 40 + int((i / total_paras) * 50)
                update_progress(task_id, "processing", prog, f"Traduciendo párrafo {i} de {total_paras}...")

    print("Translating tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        translated_text = translate_text(client, text, direction)
                        para.text = translated_text

    print(f"Saving translated document to '{output_path}'...")
    try:
        doc.save(output_path)
        print("Done!")
    except Exception as e:
        print(f"Error saving DOCX: {e}")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description="Translate academic articles (PDF or DOCX) for Mathematics Education.")
    parser.add_argument("input_file", help="Path to the input PDF or DOCX file.")
    parser.add_argument("--direction", choices=["en2es", "es2en"], default="en2es", 
                        help="Translation direction: 'en2es' (English to Spanish) or 'es2en' (Spanish to English). Default is en2es.")
    parser.add_argument("--output", help="Optional output file path. If not provided, a suffix will be appended.")
    
    args = parser.parse_args()
    
    load_dotenv()
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("Error: GEMINI_API_KEY environment variable not set. Please add it to the .env file.")
        sys.exit(1)

    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        print(f"Error initializing Gemini Client: {e}")
        sys.exit(1)

    input_path = args.input_file
    if not os.path.exists(input_path):
        print(f"Error: File '{input_path}' not found.")
        sys.exit(1)

    base_name, ext = os.path.splitext(input_path)
    ext = ext.lower()
    
    suffix = "_ES" if args.direction == "en2es" else "_EN"
    
    if args.output:
        output_path = args.output
    else:
        output_path = f"{base_name}{suffix}.docx"

    # If it's a PDF, first convert to DOCX
    if ext == ".pdf":
        temp_docx_path = f"{base_name}_temp_conversion.docx"
        convert_pdf_to_docx(input_path, temp_docx_path)
        process_docx(temp_docx_path, output_path, client, args.direction)
        # Clean up temporary DOCX
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
    elif ext == ".docx":
        process_docx(input_path, output_path, client, args.direction)
    else:
        print(f"Error: Unsupported file extension '{ext}'. Only .pdf and .docx are supported.")
        sys.exit(1)

if __name__ == "__main__":
    main()
